from __future__ import annotations

import shutil
import subprocess
import textwrap
from pathlib import Path

import markdown
from bs4 import BeautifulSoup
from docx import Document
from pdf2docx import Converter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen.canvas import Canvas
from sqlalchemy.orm import Session

from ..config import settings
from ..database import SessionLocal
from ..models import FileKind, Job, JobStatus, StoredFile
from .cache import mirror_job


def create_job(db: Session, file_id: int, kind: str) -> Job:
    job = Job(file_id=file_id, kind=kind, status=JobStatus.queued)
    db.add(job)
    db.commit()
    db.refresh(job)
    mirror_job(job)
    return job


def run_job(job_id: int) -> None:
    db = SessionLocal()
    try:
        job = db.get(Job, job_id)
        if not job:
            return
        job.status = JobStatus.running
        db.commit()
        mirror_job(job)

        file = db.get(StoredFile, job.file_id)
        if not file:
            raise RuntimeError("Source file was not found.")

        if job.kind == "markdown-to-pdf":
            output = markdown_to_pdf(file)
        elif job.kind == "markdown-to-docx":
            output = markdown_to_docx(file)
        elif job.kind == "pdf-to-docx":
            output = pdf_to_docx(file)
        else:
            raise RuntimeError(f"Unsupported job kind: {job.kind}")

        job.output_path = str(output)
        job.status = JobStatus.completed
        job.error = None
        mirror_job(job)
    except Exception as exc:
        job = db.get(Job, job_id)
        if job:
            job.status = JobStatus.failed
            job.error = str(exc)
            mirror_job(job)
    finally:
        db.commit()
        db.close()


def _output_path(source: StoredFile, suffix: str) -> Path:
    settings.exports_dir.mkdir(parents=True, exist_ok=True)
    stem = Path(source.name).stem or "document"
    return settings.exports_dir / f"{source.id}-{stem}{suffix}"


def _run_pandoc(source: Path, output: Path, target_format: str) -> bool:
    if not shutil.which("pandoc"):
        return False
    subprocess.run(
        ["pandoc", str(source), "-o", str(output), "--standalone"],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    return output.exists()


def markdown_to_docx(file: StoredFile) -> Path:
    if file.kind != FileKind.markdown:
        raise RuntimeError("Only Markdown files can be exported to Word.")
    source = Path(file.path)
    output = _output_path(file, ".docx")

    if _run_pandoc(source, output, "docx"):
        return output

    html = markdown.markdown(source.read_text(encoding="utf-8"), extensions=["extra", "tables", "fenced_code"])
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "pre", "blockquote"]):
        text = element.get_text("\n").strip()
        if not text:
            continue
        if element.name and element.name.startswith("h"):
            doc.add_heading(text, level=min(int(element.name[1]), 4))
        elif element.name == "li":
            doc.add_paragraph(text, style="List Bullet")
        elif element.name == "pre":
            doc.add_paragraph(text, style="Intense Quote")
        else:
            doc.add_paragraph(text)
    doc.save(output)
    return output


def markdown_to_pdf(file: StoredFile) -> Path:
    if file.kind != FileKind.markdown:
        raise RuntimeError("Only Markdown files can be exported to PDF.")
    source = Path(file.path)
    output = _output_path(file, ".pdf")

    if _run_pandoc(source, output, "pdf"):
        return output

    html = markdown.markdown(source.read_text(encoding="utf-8"), extensions=["extra", "tables", "fenced_code"])
    text = BeautifulSoup(html, "html.parser").get_text("\n")
    canvas = Canvas(str(output), pagesize=A4)
    width, height = A4
    x = 48
    y = height - 56
    canvas.setFont("Helvetica", 10)

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        wrapped = textwrap.wrap(line, width=92) or [""]
        for part in wrapped:
            if y < 48:
                canvas.showPage()
                canvas.setFont("Helvetica", 10)
                y = height - 56
            canvas.drawString(x, y, part.encode("latin-1", errors="replace").decode("latin-1"))
            y -= 14
    canvas.save()
    return output


def pdf_to_docx(file: StoredFile) -> Path:
    if file.kind != FileKind.pdf:
        raise RuntimeError("Only PDF files can be converted to Word.")
    output = _output_path(file, ".docx")
    converter = Converter(file.path)
    try:
        converter.convert(str(output), start=0, end=None)
    finally:
        converter.close()
    return output
